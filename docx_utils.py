# import the packages
import os, warnings, logging
from typing import List, Dict, Union, Optional

from docx import Document

logging.basicConfig(
    level = logging.INFO,
    format = '[%(levelname)s] %(asctime)s — %(message)s',
    datefmt = '%Y-%m-%d %H:%M:%S'
)


# define the function of replacing the template placeholder in tables with item dictionary information
def populate_docx_table(
        item_dict: Dict[str, str],
        docx_template_path: str,
        new_docx_path: str) -> None:
    """
    Populate a DOCX file table using placeholder keys and a data dictionary.

    Parameters:
        item_dict (dict): Dictionary with keys matching placeholders in the template.
        docx_template_path (str): Path to the input template.
        new_docx_path (str): Path to save the updated document.

    Returns:
        None

    Raises:
        TypeError: If the input types are invalid.
        FileNotFoundError: If the template file does not exist or cannot be opened.
    """

    # check the errors for file type
    if isinstance(item_dict, dict) == False:
        raise TypeError("item_dict should be a dictionary.")
    elif not docx_template_path.endswith('.docx'):
        raise TypeError("docx_template_path should be a DOCX file.")
    elif not new_docx_path.endswith('.docx'):
        raise TypeError("new_docx_path should be a DOCX file.")

    # open the template DOCX
    try:
        doc = Document(docx_template_path)
    except:
        raise FileNotFoundError("Error: template file not found or is not a valid DOCX file.")

    # replace the placeholder in the DOCX for all item_dict information
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text in item_dict.keys():
                            run.text = item_dict[run.text]

    # save the DOCX to the new_docx_path; log reminder
    doc.save(new_docx_path)
    logging.info(f"DOCX generated — Path: {new_docx_path}")

#---

# define the function of replacing the template placeholder in paragraphs with item dictionary information
def populate_docx_paragraph(
        item_dict: Dict[str, str],
        docx_template_path: str,
        new_docx_path: str) -> None:
    """
    Populate a DOCX file paragraph using placeholder keys and a data dictionary.

    Parameters:
        item_dict (dict): Dictionary with keys matching placeholders in the template.
        docx_template_path (str): Path to the input template.
        new_docx_path (str): Path to save the updated document.

    Returns:
        None

    Raises:
        TypeError: If the input types are invalid.
        FileNotFoundError: If the template file does not exist or cannot be opened.
    """

    # check the errors for file type
    if isinstance(item_dict, dict) == False:
        raise TypeError("item_dict should be a dictionary.")
    elif not docx_template_path.endswith('.docx'):
        raise TypeError("docx_template_path should be a DOCX file.")
    elif not new_docx_path.endswith('.docx'):
        raise TypeError("new_docx_path should be a DOCX file.")

    # open the template DOCX
    try:
        doc = Document(docx_template_path)
    except:
        raise FileNotFoundError("Error: template file not found or is not a valid DOCX file.")

    # replace the placeholder in the DOCX for all the item_dict information
    for para in doc.paragraphs:
        for run in para.runs:
            for key in item_dict.keys():
                if key in run.text:
                    run.text = run.text.replace(key, item_dict[key])

    # save the DOCX to the new_docx_path; log reminder
    doc.save(new_docx_path)
    logging.info(f"DOCX generated — Path: {new_docx_path}")
