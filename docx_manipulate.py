# import the packages
import os, warnings, logging
from typing import List, Dict, Union, Optional

from docx import Document
from docx2pdf import convert
from PyPDF2 import PdfMerger

logging.basicConfig(
    level = logging.INFO,
    format = '[%(levelname)s] %(asctime)s — %(message)s',
    datefmt = '%Y-%m-%d %H:%M:%S'
)


# define the function of replacing the template placeholder in tables with item dictionary information
def populate_docx_table(
        item_dict: Dict[str, str],
        docx_template_path: str,
        new_docx_path: str) -> None:
    """
    Populate a DOCX file table using placeholder keys and a data dictionary.

    Parameters:
        item_dict (dict): Dictionary with keys matching placeholders in the template.
        docx_template_path (str): Path to the input template.
        new_docx_path (str): Path to save the updated document.

    Returns:
        None

    Raises:
        TypeError: If the input types are invalid.
        FileNotFoundError: If the template file does not exist or cannot be opened.
    """

    # check the errors for file type
    if isinstance(item_dict, dict) == False:
        raise TypeError("item_dict should be a dictionary.")
    elif not docx_template_path.endswith('.docx'):
        raise TypeError("docx_template_path should be a DOCX file.")
    elif not new_docx_path.endswith('.docx'):
        raise TypeError("new_docx_path should be a DOCX file.")

    # open the template DOCX
    try:
        doc = Document(docx_template_path)
    except:
        raise FileNotFoundError("Error: template file not found or is not a valid DOCX file.")

    # replace the placeholder in the DOCX for all item_dict information
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text in item_dict.keys():
                            run.text = item_dict[run.text]

    # save the DOCX to the new_docx_path; log reminder
    doc.save(new_docx_path)
    logging.info(f"DOCX generated — Path: {new_docx_path}")

#---

# define the function of replacing the template placeholder in paragraphs with item dictionary information
def populate_docx_paragraph(
        item_dict: Dict[str, str],
        docx_template_path: str,
        new_docx_path: str) -> None:
    """
    Populate a DOCX file paragraph using placeholder keys and a data dictionary.

    Parameters:
        item_dict (dict): Dictionary with keys matching placeholders in the template.
        docx_template_path (str): Path to the input template.
        new_docx_path (str): Path to save the updated document.

    Returns:
        None

    Raises:
        TypeError: If the input types are invalid.
        FileNotFoundError: If the template file does not exist or cannot be opened.
    """

    # check the errors for file type
    if isinstance(item_dict, dict) == False:
        raise TypeError("item_dict should be a dictionary.")
    elif not docx_template_path.endswith('.docx'):
        raise TypeError("docx_template_path should be a DOCX file.")
    elif not new_docx_path.endswith('.docx'):
        raise TypeError("new_docx_path should be a DOCX file.")

    # open the template DOCX
    try:
        doc = Document(docx_template_path)
    except:
        raise FileNotFoundError("Error: template file not found or is not a valid DOCX file.")

    # replace the placeholder in the DOCX for all the item_dict information
    for para in doc.paragraphs:
        for run in para.runs:
            for key in item_dict.keys():
                if key in run.text:
                    run.text = run.text.replace(key, item_dict[key])

    # save the DOCX to the new_docx_path; log reminder
    doc.save(new_docx_path)
    logging.info(f"DOCX generated — Path: {new_docx_path}")

#---

def convert_docx_pdf(
        docx_path: str,
        keep: bool = True) -> None:
    """
    Convert a DOCX file to a PDF file.

    Parameters:
        docx_path (str): Path to the input DOCX file.
        keep (bool, optional): Whether to keep the original DOCX file after conversion. Defaults to True.

    Returns:
        None

    Raises:
        TypeError: If the input path is not a DOCX file.
    """

    # check the errors for file type
    if not docx_path.endswith('.docx'):
        raise TypeError("docx_path should be a DOCX file.")

    # convert the DOCX to PDF
    convert(docx_path)

    # remove the docx if needed; log reminder
    if keep == False:
        os.remove(docx_path)
        logging.info(f"PDF generated (original DOCX removed) — Path: {docx_path}")
    else:
        logging.info(f"PDF generated (original DOCX kept) — Path: {docx_path}")

#---

def merge_pdfs(
        pdf_list: Union[List[str], str],
        output_path: str) -> None:
    """
    Merge multiple PDF files into a single PDF document.

    Parameters:
        pdf_list (list or str): A list of PDF file paths or a folder path containing PDF files.
        output_path (str): The path where the merged PDF will be saved. Must end with '.pdf'.

    Returns:
        None

    Raises:
        TypeError: If the input is not a list of PDF paths or a valid folder path,
                   or if any file in the original PDF list is not a PDF,
                   or if the output file path does not end with '.pdf',
                   or if the final PDF list for combination does not contain any PDF file.
        UserWarning: If the folder contains any non-PDF files.
    """

    # check the errors for file type
    if not output_path.endswith('.pdf'):
        raise TypeError("output_path should be a PDF file.")

    if isinstance(pdf_list, list):
        for file in pdf_list:
            if not file.endswith('.pdf'):
                raise TypeError(f"{file} is not a PDF file.")

    elif os.path.isdir(pdf_list):
        folder_path = pdf_list
        pdf_list = os.listdir(pdf_list)
        non_pdf_list = [file for file in pdf_list if not file.endswith('.pdf')]
        pdf_list = [folder_path + '/' + file for file in pdf_list if file.endswith('.pdf')]

        if len(non_pdf_list) >= 1:
            warnings.warn(f"{len(non_pdf_list)} non-PDF file(s) detected in the folder. They will be ignored during merging.")

    else:
        raise TypeError(f"{pdf_list} should be a list of PDF files or a valid folder path.")

    # check the length of the pdf_list
    if len(pdf_list) == 0:
        raise TypeError("There is no PDF file in the file list.")
    else:
        logging.info("Merging the following PDF files:")
        for file in pdf_list:
            logging.info(f"  - {file}")

    # merge the PDF files
    merger = PdfMerger()

    for pdf in pdf_list:
        merger.append(pdf)

    # save the merged PDF file to the output path; log reminder
    merger.write(output_path)
    merger.close()
    logging.info(f"PDF merged — Path: {output_path}")
